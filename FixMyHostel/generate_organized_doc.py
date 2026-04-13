import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

root = r'c:\xampp\htdocs\FixMyHostel'
img_dir = os.path.join(root, 'Screenshots')
output_path = os.path.join(root, 'FixMyHostel_Documentation.docx')

doc = Document()

# Add title
title = doc.add_heading('FixMyHostel - Complete Documentation', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('A comprehensive guide to the hostel complaint management system for all user roles.')

# ===== SECTION 1: LOGIN =====
doc.add_heading('1. Main Login Page', level=1)
doc.add_paragraph(
    'The main login page allows Students, Maintenance Staff, and Administrators to access the system using their credentials. '
    'All users login through the same page and are automatically routed to their respective dashboards based on their role.'
)
if os.path.exists(os.path.join(img_dir, 'Login.png')):
    doc.add_picture(os.path.join(img_dir, 'Login.png'), width=Inches(6))
    doc.add_paragraph('Main Login Page', style='Caption')

# ===== SECTION 2: STUDENT ROLE =====
doc.add_page_break()
doc.add_heading('2. Student Role & Workflows', level=1)

student_desc = doc.add_paragraph()
student_desc.add_run('Student Capabilities:\n').bold = True
student_desc.add_run(
    '• View personal dashboard with complaint statistics\n'
    '• Submit new complaints about hostel issues\n'
    '• Track the status of submitted complaints\n'
    '• Edit pending complaints\n'
    '• Delete pending complaints\n'
    '• View complaint details and responses'
)

doc.add_heading('2.1 Student Login', level=2)
if os.path.exists(os.path.join(img_dir, 'Login_Student.png')):
    doc.add_picture(os.path.join(img_dir, 'Login_Student.png'), width=Inches(6))

doc.add_heading('2.2 Student Dashboard', level=2)
doc.add_paragraph(
    'The student dashboard displays an overview of all complaints with status breakdown (Pending, In Progress, Resolved). '
    'Shows Complaint ID, Title, Category, Description, Status, and Created Date.'
)
if os.path.exists(os.path.join(img_dir, 'Student_dashboard.png')):
    doc.add_picture(os.path.join(img_dir, 'Student_dashboard.png'), width=Inches(6))
    doc.add_paragraph('Student Dashboard', style='Caption')

doc.add_heading('2.3 Submit Complaint', level=2)
doc.add_paragraph(
    'Students can submit complaints by filling in the complaint form with details like title, category, priority level, description, and images.'
)
if os.path.exists(os.path.join(img_dir, 'student_complaint_submit.png')):
    doc.add_picture(os.path.join(img_dir, 'student_complaint_submit.png'), width=Inches(6))
    doc.add_paragraph('Submit Complaint Form', style='Caption')

doc.add_heading('2.4 Track Complaints', level=2)
doc.add_paragraph(
    'Students can view all their complaints with real-time status updates. They can see Student Name, Complaint Title, Category, Priority, Status, and Created Date. '
    'Pending complaints can be edited or deleted.'
)
if os.path.exists(os.path.join(img_dir, 'student_track_complaints.png')):
    doc.add_picture(os.path.join(img_dir, 'student_track_complaints.png'), width=Inches(6))
    doc.add_paragraph('Student Complaints List', style='Caption')

# ===== SECTION 3: MAINTENANCE ASSOCIATE =====
doc.add_page_break()
doc.add_heading('3. Maintenance Associate Role & Workflows', level=1)

maint_desc = doc.add_paragraph()
maint_desc.add_run('Maintenance Associate Capabilities:\n').bold = True
maint_desc.add_run(
    '• View all complaints assigned to Maintenance department\n'
    '• Pick up pending complaints (change status to "In Progress")\n'
    '• View complaint details and student information\n'
    '• Track complaint progress\n'
    '• View complaint priority and category'
)

doc.add_heading('3.1 Maintenance Associate Login', level=2)
if os.path.exists(os.path.join(img_dir, 'Maintenance_Assoc_Login.png')):
    doc.add_picture(os.path.join(img_dir, 'Maintenance_Assoc_Login.png'), width=Inches(6))

doc.add_heading('3.2 Maintenance Associate Dashboard', level=2)
doc.add_paragraph(
    'The dashboard shows all maintenance complaints with Student Name, Complaint Title, Category, Priority, Status, and Action options. '
    'Associates can pick up pending complaints to start working on them.'
)
if os.path.exists(os.path.join(img_dir, 'Maintenance_Assoc_Dashboard.png')):
    doc.add_picture(os.path.join(img_dir, 'Maintenance_Assoc_Dashboard.png'), width=Inches(6))
    doc.add_paragraph('Maintenance Associate Dashboard', style='Caption')

doc.add_heading('3.3 Pick Up Complaint', level=2)
doc.add_paragraph(
    'Associates can pick up a pending complaint by selecting it and changing the status to "In Progress". This indicates they are actively working on resolving the issue.'
)
if os.path.exists(os.path.join(img_dir, 'Maintenance_Assoc_Pickup_Complaint.png')):
    doc.add_picture(os.path.join(img_dir, 'Maintenance_Assoc_Pickup_Complaint.png'), width=Inches(6))
    doc.add_paragraph('Pick Up Complaint', style='Caption')

# ===== SECTION 4: MAINTENANCE SUPERVISOR =====
doc.add_page_break()
doc.add_heading('4. Maintenance Supervisor Role & Workflows', level=1)

super_desc = doc.add_paragraph()
super_desc.add_run('Maintenance Supervisor Capabilities:\n').bold = True
super_desc.add_run(
    '• View all maintenance complaints\n'
    '• Review in-progress complaints\n'
    '• Resolve complaints (change status to "Resolved")\n'
    '• Verify completed work by associates\n'
    '• View complaint details and history\n'
    '• Track complaint timelines'
)

doc.add_heading('4.1 Maintenance Supervisor Login', level=2)
if os.path.exists(os.path.join(img_dir, 'Supervisor_Login.png')):
    doc.add_picture(os.path.join(img_dir, 'Supervisor_Login.png'), width=Inches(6))

doc.add_heading('4.2 Supervisor Dashboard', level=2)
doc.add_paragraph(
    'Supervisors see all maintenance complaints and can only resolve complaints that are in "In Progress" status. '
    'This ensures quality control before marking complaints as resolved.'
)
if os.path.exists(os.path.join(img_dir, 'Supervisor_Dashboard.png')):
    doc.add_picture(os.path.join(img_dir, 'Supervisor_Dashboard.png'), width=Inches(6))
    doc.add_paragraph('Supervisor Dashboard', style='Caption')

doc.add_heading('4.3 Resolve Complaint', level=2)
doc.add_paragraph(
    'The supervisor can verify the work done by the associate and mark the complaint as "Resolved" after confirming it meets standards.'
)
if os.path.exists(os.path.join(img_dir, 'supervisor_resolve_complaint_after_verify.png')):
    doc.add_picture(os.path.join(img_dir, 'supervisor_resolve_complaint_after_verify.png'), width=Inches(6))
    doc.add_paragraph('Resolve Complaint', style='Caption')

# ===== SECTION 5: ADMIN =====
doc.add_page_break()
doc.add_heading('5. Administrator Role & Workflows', level=1)

admin_desc = doc.add_paragraph()
admin_desc.add_run('Administrator Capabilities:\n').bold = True
admin_desc.add_run(
    '• Complete system overview with statistics\n'
    '• View all complaints across all departments\n'
    '• Manage user accounts (edit, remove, reset passwords)\n'
    '• Change complaint status at any stage (Pending → In Progress → Resolved → Closed)\n'
    '• View detailed complaint information\n'
    '• Monitor system usage and statistics\n'
    '• Manage multiple departments'
)

doc.add_heading('5.1 Admin Login', level=2)
if os.path.exists(os.path.join(img_dir, 'Admin_Login.png')):
    doc.add_picture(os.path.join(img_dir, 'Admin_Login.png'), width=Inches(6))

doc.add_heading('5.2 Admin Dashboard', level=2)
doc.add_paragraph(
    'The admin dashboard displays comprehensive statistics including Total Complaints, Pending, In Progress, Resolved, Total Users, Students, and Maintenance Staff. '
    'All complaints from all departments are visible with full details.'
)
if os.path.exists(os.path.join(img_dir, 'Admin_dashboard.png')):
    doc.add_picture(os.path.join(img_dir, 'Admin_dashboard.png'), width=Inches(6))
    doc.add_paragraph('Admin Dashboard', style='Caption')

doc.add_heading('5.3 Review Complaints', level=2)
doc.add_paragraph(
    'Admins can review all complaints from all departments and have the ability to change their status at any point in the workflow.'
)
if os.path.exists(os.path.join(img_dir, 'Admin_review_Complaints.png')):
    doc.add_picture(os.path.join(img_dir, 'Admin_review_Complaints.png'), width=Inches(6))
    doc.add_paragraph('Review All Complaints', style='Caption')

doc.add_heading('5.4 Manage Users', level=2)
doc.add_paragraph(
    'Administrators can manage all user accounts in the system, including students and staff members. They can edit user details and reset passwords as needed.'
)
if os.path.exists(os.path.join(img_dir, 'Admin_Manage_Users.png')):
    doc.add_picture(os.path.join(img_dir, 'Admin_Manage_Users.png'), width=Inches(6))
    doc.add_paragraph('Manage Users', style='Caption')

doc.add_heading('5.5 Edit Student Details', level=2)
doc.add_paragraph(
    'Administrators can update student information and account settings.'
)
if os.path.exists(os.path.join(img_dir, 'Admin_Edit_Student_Details.png')):
    doc.add_picture(os.path.join(img_dir, 'Admin_Edit_Student_Details.png'), width=Inches(6))
    doc.add_paragraph('Edit Student Details', style='Caption')

# ===== SUMMARY =====
doc.add_page_break()
doc.add_heading('6. System Summary', level=1)

summary = doc.add_paragraph()
summary.add_run('User Roles Overview:\n\n').bold = True
summary.add_run(
    '1. Students: Submit and track complaints, limited editing capabilities\n\n'
    '2. Maintenance Associates: Pickup and work on complaints, report to supervisors\n\n'
    '3. Maintenance Supervisors: Review and resolve complaints from associates\n\n'
    '4. Administrators: Full system control, manage users and monitor all complaints\n\n'
    '\nComplaint Status Flow:\n'
    'Pending → In Progress → Resolved → Closed\n\n'
    'Each role has specific permissions at each stage to ensure proper workflow management.'
)

doc.save(output_path)
print(f"Document created: {output_path}")
