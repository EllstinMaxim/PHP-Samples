import os
from docx import Document
from docx.shared import Inches

root = r'c:\xampp\htdocs\FixMyHostel'
img_dir = os.path.join(root, 'Screenshots')
output_path = os.path.join(root, 'FixMyHostel_Screenshots_Documentation.docx')

doc = Document()
doc.add_heading('FixMyHostel Screenshot Documentation', level=1)
doc.add_paragraph('This document contains all screenshots from the Screenshots folder, with filenames as captions.')

files = sorted([f for f in os.listdir(img_dir) if os.path.isfile(os.path.join(img_dir, f))])
for fname in files:
    path = os.path.join(img_dir, fname)
    doc.add_heading(fname, level=2)
    try:
        doc.add_picture(path, width=Inches(6))
    except Exception as e:
        doc.add_paragraph(f'Unable to insert image: {e}')
    doc.add_paragraph(fname, style='Caption')

doc.save(output_path)
print(output_path)
